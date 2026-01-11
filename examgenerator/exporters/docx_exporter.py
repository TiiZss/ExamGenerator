"""
DOCX exporter for exams.
"""

import os
import random
from datetime import datetime
from typing import List, Dict, Optional, Any, TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document as DocumentType
else:
    DocumentType = Any

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None  # type: ignore
    Pt = None  # type: ignore
    Inches = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    WD_STYLE_TYPE = None  # type: ignore


def create_docx_document(title: str, template_path: Optional[str] = None) -> Any:
    """Create a new DOCX document with custom styles or from template.
    
    Args:
        title: Document title
        template_path: Optional path to template file
        
    Returns:
        Document object
        
    Raises:
        ImportError: If python-docx is not installed
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx no está instalado. Ejecuta: uv add python-docx")
    
    if template_path and os.path.exists(template_path):
        # Load from template
        doc = Document(template_path)
        print(f"Usando plantilla: {template_path}")
    else:
        # Create new document with default styles
        doc = Document()
        
        # Configure document styles
        styles = doc.styles
        
        # Title style
        if 'Custom Title' not in [style.name for style in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Question style
        if 'Question' not in [style.name for style in styles]:
            question_style = styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
            question_style.font.size = Pt(11)
            question_style.font.bold = True
            question_style.paragraph_format.space_before = Pt(6)
            question_style.paragraph_format.space_after = Pt(3)
        
        # Option style
        if 'Option' not in [style.name for style in styles]:
            option_style = styles.add_style('Option', WD_STYLE_TYPE.PARAGRAPH)
            option_style.font.size = Pt(10)
            option_style.paragraph_format.left_indent = Inches(0.5)
            option_style.paragraph_format.space_after = Pt(2)
    
    return doc


def replace_placeholders(doc: Any, exam_prefix: str, exam_number: int, num_questions: int):
    """Replace placeholders in the template document.
    
    Args:
        doc: Document object
        exam_prefix: Exam prefix (e.g., "Parcial")
        exam_number: Exam number
        num_questions: Number of questions in exam
    """
    from ..core.time_calculator import calculate_exam_time
    
    now = datetime.now()
    
    # Spanish month names
    months_es = {
        1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
        5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
        9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
    }
    
    # Calculate exam time
    exam_time = calculate_exam_time(num_questions)
    exam_time_minutes = int(num_questions * 1.0)  # Default: 1 minute per question
    
    replacements = {
        '{{EXAM_NAME}}': f"{exam_prefix}",
        '{{EXAM_NUMBER}}': str(exam_number),
        '{{EXAM_TITLE}}': f"EXAMEN {exam_prefix} {exam_number}",
        '{{DATE}}': now.strftime("%d/%m/%Y"),
        '{{FULL_DATE}}': now.strftime("%d de %B de %Y"),
        '{{DAY}}': str(now.day),
        '{{MONTH}}': months_es[now.month],
        '{{MONTH_NUM}}': f"{now.month:02d}",
        '{{YEAR}}': str(now.year),
        '{{COURSE}}': exam_prefix,
        '{{TODAY}}': now.strftime("%d/%m/%Y"),
        '{{MONTH_YEAR}}': f"{months_es[now.month]} {now.year}",
        '{{NUM_QUESTIONS}}': str(num_questions),
        '{{EXAM_TIME}}': exam_time,
        '{{EXAM_TIME_MINUTES}}': str(exam_time_minutes),
        '{{TIME}}': exam_time,
        '{{DURATION}}': exam_time
    }
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)
    
    # Replace in tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, replacement)


def find_content_insertion_point(doc: Any) -> Optional[int]:
    """Find where to insert exam content in the template.
    
    Args:
        doc: Document object
        
    Returns:
        Paragraph index for insertion, or None if not found
    """
    content_markers = ['{{CONTENT}}', '{{QUESTIONS}}', '{{EXAM_CONTENT}}']
    
    for i, paragraph in enumerate(doc.paragraphs):
        for marker in content_markers:
            if marker in paragraph.text:
                # Remove the marker
                paragraph.text = ""
                return i
    
    # If no marker found, insert after the first paragraph or at the end
    return len(doc.paragraphs) if len(doc.paragraphs) > 0 else 0


def create_exam_docx(
    exam_prefix: str,
    exam_number: int,
    selected_questions: List[Dict],
    output_dir: str,
    template_path: Optional[str] = None
) -> None:
    """Save exam as DOCX file.
    
    Args:
        exam_prefix: Exam prefix (e.g., "Parcial", "Final")
        exam_number: Exam number
        selected_questions: List of question dictionaries with shuffled options
        output_dir: Output directory path
        template_path: Optional path to DOCX template
        
    Raises:
        ImportError: If python-docx is not installed
    """
    if not DOCX_AVAILABLE:
        print("Error: Para exportar a DOCX necesitas instalar python-docx:")
        print("uv add python-docx")
        return
    
    option_letters = ['A', 'B', 'C', 'D']
    num_questions = len(selected_questions)
    
    # Create exam document
    exam_doc = create_docx_document(f"EXAMEN {exam_prefix} {exam_number}", template_path)
    
    # Replace placeholders in template
    replace_placeholders(exam_doc, exam_prefix, exam_number, num_questions)
    
    # Find insertion point for content
    insertion_point = find_content_insertion_point(exam_doc)
    
    # If no template or no insertion point found, add title
    if not template_path or insertion_point == 0:
        title_para = exam_doc.add_paragraph(f"EXAMEN {exam_prefix} {exam_number}")
        try:
            title_para.style = 'Custom Title'
        except KeyError:
            # Style doesn't exist, apply manual formatting
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.size = Pt(16)
                run.font.bold = True
    
    # Add questions
    for i, q_data in enumerate(selected_questions, 1):
        question_text = q_data["question"]
        shuffled_options = q_data["options"]  # Already shuffled in exam_generator
        
        # Add question
        question_para = exam_doc.add_paragraph(f"{i}. {question_text}")
        try:
            question_para.style = 'Question'
        except KeyError:
            # Style doesn't exist, apply manual formatting
            for run in question_para.runs:
                run.font.bold = True
                run.font.size = Pt(11)
        
        # Add options
        for j, option in enumerate(shuffled_options):
            option_para = exam_doc.add_paragraph(f"{option_letters[j]}) {option}")
            try:
                option_para.style = 'Option'
            except KeyError:
                # Style doesn't exist, apply manual formatting
                option_para.paragraph_format.left_indent = Inches(0.5)
                for run in option_para.runs:
                    run.font.size = Pt(10)
        
        # Add some space after each question
        exam_doc.add_paragraph()
    
    # Save exam document
    filename = os.path.join(output_dir, f"examen_{exam_prefix}_{exam_number}.docx")
    exam_doc.save(filename)
    print(f"Examen DOCX creado: {filename}")
